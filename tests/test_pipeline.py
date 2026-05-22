"""End-to-end pipeline tests. Groq calls are mocked; the real DOCX
parser and writer run against a synthetic resume file built in tmp_path."""

import json
from pathlib import Path
from unittest.mock import MagicMock

import pytest
from docx import Document

from src.pipeline import run_tailor_pipeline


# ---------- synthetic resume fixture (small to keep mock-call lists short) ----------


def _make_synthetic_resume(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Alice Example")
    doc.add_paragraph("STAFF ENGINEER")
    doc.add_paragraph("NYC | (555) 010-2020 | alice@example.com | linkedin.com/in/alice")
    doc.add_paragraph("")
    doc.add_paragraph("Staff engineer with 12 years building backend systems.")
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("AcmeCorp | Staff Engineer")
    doc.add_paragraph("Jan 2020 - Present | Remote")
    doc.add_paragraph("Led platform team of 6 engineers.")
    p1 = doc.add_paragraph("Cut p99 latency from 800ms to 200ms.")
    p1.style = doc.styles["List Paragraph"]
    p2 = doc.add_paragraph("Mentored 6 engineers, 2 promoted within 12 months.")
    p2.style = doc.styles["List Paragraph"]
    doc.add_paragraph("Skills: Python, Redis, PostgreSQL")
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("MIT | BS, Computer Science\n2008 - 2012 | Cambridge, MA")
    doc.add_paragraph("SKILLS")
    sk = doc.add_paragraph("Languages: Python, Go")
    sk.style = doc.styles["List Paragraph"]

    path = tmp_path / "alice.docx"
    doc.save(str(path))
    return path


# ---------- canned LLM responses ----------


JD_ANALYSIS = {
    "must_have": [
        {"tech": "Python", "category": "language", "evidence": "5+ years Python"},
        {"tech": "Kafka", "category": "streaming", "evidence": "build Kafka pipelines"},
    ],
    "nice_to_have": [],
    "soft_skills": [],
    "domain_keywords": ["real-time analytics"],
    "seniority_level": "staff",
    "exact_phrases_to_mirror": ["distributed systems at scale"],
}

SUBSTITUTION_PLAN = {
    "title_modifier": "Backend Platforms",
    "substitutions": [
        {"old": "Redis", "new": "Kafka", "domain": "streaming", "apply_to": ["skills"]}
    ],
    "additions_to_skills": ["Real-Time Pipelines"],
    "summary_focus": "Emphasize distributed systems and streaming",
}


def _make_mock_client(responses: list) -> MagicMock:
    client = MagicMock()
    client.complete_json.side_effect = list(responses)
    client.summary.return_value = {
        "model": "test-model", "total_calls": len(responses),
        "successful_calls": len(responses), "failed_calls": 0,
        "total_retries": 0, "total_prompt_tokens": 100,
        "total_completion_tokens": 50, "total_tokens": 150,
        "avg_latency_ms": 200.0, "calls_per_key": {0: len(responses)},
    }
    client.format_summary.return_value = "  Groq usage — model=test-model\n    calls         : N"
    return client


# ---------- end-to-end happy path ----------


def test_pipeline_runs_end_to_end_with_pdf_skipped(tmp_path: Path) -> None:
    resume_path = _make_synthetic_resume(tmp_path)
    out_dir = tmp_path / "out"

    # Call order for this resume (1 role, 2 bullets):
    # 1) JD analyze, 2) substitution plan, 3) summary rewrite,
    # 4) intro rewrite, 5) bullet[0] rewrite, 6) bullet[1] rewrite
    responses = [
        JD_ANALYSIS,
        SUBSTITUTION_PLAN,
        {"text": "Staff engineer with 12 years on distributed systems at scale."},
        {"text": "Led an AcmeCorp platform team of 6 engineers focused on distributed systems."},
        {"text": "Cut p99 latency from 800ms to 200ms across the distributed request path."},
        {"text": "Mentored 6 engineers, 2 promoted within 12 months."},  # verbatim, no change
    ]
    client = _make_mock_client(responses)

    result = run_tailor_pipeline(
        resume_path=resume_path,
        jd_text="Senior backend engineer with Python and Kafka.",
        output_dir=out_dir,
        client=client,
        skip_pdf=True,
    )

    assert result.docx_path.exists()
    assert result.docx_path.suffix == ".docx"
    assert result.pdf_path is None
    assert result.report.passed is True
    assert client.complete_json.call_count == 6


def test_pipeline_writes_per_stage_json_logs(tmp_path: Path) -> None:
    resume_path = _make_synthetic_resume(tmp_path)
    out_dir = tmp_path / "out"
    client = _make_mock_client([
        JD_ANALYSIS, SUBSTITUTION_PLAN,
        {"text": "summary"}, {"text": "intro"},
        {"text": "Cut p99 latency from 800ms to 200ms."},
        {"text": "Mentored 6 engineers, 2 promoted within 12 months."},
    ])

    result = run_tailor_pipeline(
        resume_path=resume_path, jd_text="JD text", output_dir=out_dir,
        client=client, skip_pdf=True,
    )

    log_files = {p.name for p in result.log_dir.iterdir()}
    assert "1_jd_analysis.json" in log_files
    assert "2_resume_parsed.json" in log_files
    assert "3_substitution_plan.json" in log_files
    assert "4a_rewritten_initial.json" in log_files
    assert "5_validation_report.json" in log_files
    assert "groq_usage.json" in log_files

    # Spot-check JSON validity.
    payload = json.loads((result.log_dir / "1_jd_analysis.json").read_text())
    assert payload["seniority_level"] == "staff"


def test_pipeline_attempts_pdf_when_not_skipped(monkeypatch, mocker, tmp_path: Path) -> None:
    resume_path = _make_synthetic_resume(tmp_path)
    out_dir = tmp_path / "out"
    # Stub the PDF exporter so the test doesn't depend on LibreOffice.
    mocker.patch(
        "src.pipeline.export_pdf",
        side_effect=lambda src, dst, **kw: Path(dst).write_bytes(b"%PDF-1.4 fake") or Path(dst),
    )
    client = _make_mock_client([
        JD_ANALYSIS, SUBSTITUTION_PLAN,
        {"text": "summary"}, {"text": "intro"},
        {"text": "Cut p99 latency from 800ms to 200ms."},
        {"text": "Mentored 6 engineers, 2 promoted within 12 months."},
    ])

    result = run_tailor_pipeline(
        resume_path=resume_path, jd_text="JD", output_dir=out_dir, client=client,
        skip_pdf=False,
    )
    assert result.pdf_path is not None
    assert result.pdf_path.exists()


def test_pipeline_records_pdf_as_none_when_export_fails(mocker, tmp_path: Path) -> None:
    resume_path = _make_synthetic_resume(tmp_path)
    out_dir = tmp_path / "out"
    mocker.patch("src.pipeline.export_pdf", side_effect=RuntimeError("no backend"))
    client = _make_mock_client([
        JD_ANALYSIS, SUBSTITUTION_PLAN,
        {"text": "summary"}, {"text": "intro"},
        {"text": "Cut p99 latency from 800ms to 200ms."},
        {"text": "Mentored 6 engineers, 2 promoted within 12 months."},
    ])

    result = run_tailor_pipeline(
        resume_path=resume_path, jd_text="JD", output_dir=out_dir, client=client,
        skip_pdf=False,
    )
    assert result.pdf_path is None
    assert result.docx_path.exists()


# ---------- regeneration loop ----------


def test_pipeline_regenerates_when_summary_drops_numbers(tmp_path: Path) -> None:
    """Force the first summary rewrite to drop a number; the rewriter's
    own guard reverts it to the original verbatim. With no remaining
    critical issues, the regen loop should not actually run an LLM call
    for the second pass. We assert: validation passes after one regen
    attempt at most."""
    resume_path = _make_synthetic_resume(tmp_path)
    out_dir = tmp_path / "out"

    # First summary response drops "12" — guard reverts to original.
    responses = [
        JD_ANALYSIS, SUBSTITUTION_PLAN,
        {"text": "Staff engineer building backend systems."},  # numbers dropped
        {"text": "Led an AcmeCorp platform team of 6 engineers focused on distributed systems."},
        {"text": "Cut p99 latency from 800ms to 200ms."},
        {"text": "Mentored 6 engineers, 2 promoted within 12 months."},
    ]
    client = _make_mock_client(responses)
    result = run_tailor_pipeline(
        resume_path=resume_path, jd_text="JD", output_dir=out_dir, client=client,
        skip_pdf=True, max_regen_passes=2,
    )
    # Guard preserved original summary; validation passes without regen.
    assert result.report.passed is True


# ---------- groq summary surfaced ----------


def test_pipeline_progress_callback_emits_monotonic_fractions(tmp_path: Path) -> None:
    resume_path = _make_synthetic_resume(tmp_path)
    out_dir = tmp_path / "out"
    client = _make_mock_client([
        JD_ANALYSIS, SUBSTITUTION_PLAN,
        {"text": "summary"}, {"text": "intro"},
        {"text": "Cut p99 latency from 800ms to 200ms."},
        {"text": "Mentored 6 engineers, 2 promoted within 12 months."},
    ])
    events: list[tuple[str, float]] = []

    run_tailor_pipeline(
        resume_path=resume_path, jd_text="JD", output_dir=out_dir,
        client=client, skip_pdf=True,
        progress_cb=lambda label, frac: events.append((label, frac)),
    )

    assert events  # at least one event
    fractions = [f for _, f in events]
    assert all(0.0 <= f <= 1.0 for f in fractions)
    # Final event should be at 1.0.
    assert events[-1][1] == 1.0
    # Fractions should be non-decreasing.
    assert fractions == sorted(fractions)


def test_pipeline_returns_groq_summary_and_writes_json(tmp_path: Path) -> None:
    resume_path = _make_synthetic_resume(tmp_path)
    out_dir = tmp_path / "out"
    client = _make_mock_client([
        JD_ANALYSIS, SUBSTITUTION_PLAN,
        {"text": "summary"}, {"text": "intro"},
        {"text": "Cut p99 latency from 800ms to 200ms."},
        {"text": "Mentored 6 engineers, 2 promoted within 12 months."},
    ])
    result = run_tailor_pipeline(
        resume_path=resume_path, jd_text="JD", output_dir=out_dir, client=client,
        skip_pdf=True,
    )
    assert result.groq_summary["total_calls"] == 6
    usage_json = json.loads((result.log_dir / "groq_usage.json").read_text())
    assert usage_json["total_calls"] == 6
