"""Unit + integration tests for src.stages.resume_parser."""

from pathlib import Path

import pytest
from docx import Document

from src.stages.resume_parser import (
    ResumeParseError,
    _all_separator_positions,
    _is_contact_line,
    _split_dates_and_location,
    _split_on_separator,
    parse_resume,
)

SAMPLE_RESUME = Path(__file__).parent.parent / "profiles" / "sample" / "Katharine Berry.docx"


# ---------- helper unit tests ----------


def test_split_dates_and_location_middle_dot() -> None:
    text = "July 2018 – Present  ·  Sunnyvale, CA"
    assert _split_dates_and_location(text) == ("July 2018 – Present", "Sunnyvale, CA")


def test_split_dates_and_location_with_pipe() -> None:
    text = "2020 - 2022 | Remote"
    assert _split_dates_and_location(text) == ("2020 - 2022", "Remote")


def test_split_dates_and_location_only_range_separator() -> None:
    assert _split_dates_and_location("2010 – 2014") == ("2010 – 2014", "")


def test_split_dates_and_location_no_separator() -> None:
    assert _split_dates_and_location("2010 - 2014") == ("2010 - 2014", "")


def test_split_on_separator_first_occurrence() -> None:
    assert _split_on_separator("A – B – C", fallback=("", "")) == ("A ", " B – C")


def test_split_on_separator_fallback() -> None:
    assert _split_on_separator("no separator here", fallback=("x", "y")) == ("x", "y")


def test_all_separator_positions_sorted() -> None:
    text = "a | b – c · d"
    positions = _all_separator_positions(text)
    assert [p[0] for p in positions] == sorted(p[0] for p in positions)
    assert len(positions) == 3


def test_is_contact_line_with_pipe_and_email() -> None:
    assert _is_contact_line("Redwood City, CA | (470) 485-4029 | foo@bar.com | linkedin.com/in/x")


def test_is_contact_line_rejects_prose() -> None:
    summary = (
        "Senior Software Engineer with 11+ years of experience. Proven track record "
        "at Google, Intel, and Pebble leading developer ecosystem design."
    )
    assert not _is_contact_line(summary)


def test_is_contact_line_requires_separator() -> None:
    assert not _is_contact_line("foo@bar.com just an email")


# ---------- integration: parse the real Katharine Berry resume ----------


@pytest.fixture(scope="module")
def parsed_resume():
    if not SAMPLE_RESUME.exists():
        pytest.skip(f"sample resume not present at {SAMPLE_RESUME}")
    return parse_resume(SAMPLE_RESUME)


def test_integration_header(parsed_resume) -> None:
    h = parsed_resume.header
    assert h.name == "Katharine Berry"
    assert h.title == "SENIOR SOFTWARE ENGINEER"
    assert len(h.contact_lines) == 1
    assert "@" in h.contact_lines[0]
    assert h.indices.name_idx == 0
    assert h.indices.title_idx == 1
    assert h.indices.contact_idxs == [2]


def test_integration_summary(parsed_resume) -> None:
    s = parsed_resume.summary
    assert s.paragraph_idxs == [5]
    assert "11+ years of experience" in s.text


def test_integration_experience_count_and_first_role(parsed_resume) -> None:
    exps = parsed_resume.experience
    assert len(exps) == 4
    g = exps[0]
    assert g.company == "Google"
    assert g.title == "Senior Software Engineer"
    assert g.dates == "July 2018 – Present"
    assert g.location == "Sunnyvale, CA"
    assert len(g.bullets) == 9
    assert g.skills_line is not None and g.skills_line.startswith("Skills:")
    assert g.indices.header_idxs == [7, 8]
    assert g.indices.intro_idx == 9
    assert g.indices.bullet_idxs == [10, 11, 12, 13, 14, 15, 16, 17, 18]
    assert g.indices.skills_line_idx == 19


def test_integration_all_role_companies(parsed_resume) -> None:
    companies = [e.company for e in parsed_resume.experience]
    assert companies == ["Google", "Intel Corporation", "Pebble Technology", "Pebble Technology"]


def test_integration_education(parsed_resume) -> None:
    eds = parsed_resume.education
    assert len(eds) == 1
    e = eds[0]
    assert e.institution == "Massachusetts Institute of Technology"
    assert e.degree == "Bachelor of Science"
    assert e.field == "Computer Science"
    assert e.dates == "2010 – 2014"
    assert e.location == "Cambridge, MA"
    assert e.indices.header_idx == 50


def test_integration_skills_section(parsed_resume) -> None:
    s = parsed_resume.skills_section
    assert s.indices.header_idx == 51
    assert len(s.indices.content_idxs) == 6
    assert "Languages & Frameworks" in s.categories
    assert "Python" in s.categories["Languages & Frameworks"]


def test_integration_raw_paragraphs_count(parsed_resume) -> None:
    # Exact paragraph count from the source file — guards against silent drift.
    assert len(parsed_resume.raw_paragraphs) == 58


# ---------- synthetic resume: pipe-separator + multiple education entries ----------


def _make_synthetic_resume(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Alice Example")
    doc.add_paragraph("STAFF ENGINEER")
    doc.add_paragraph("NYC | (555) 010-2020 | alice@example.com | linkedin.com/in/alice")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Staff engineer with 12 years of experience. Built distributed systems at scale."
    )
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("AcmeCorp | Staff Engineer")
    doc.add_paragraph("Jan 2020 - Present | Remote")
    doc.add_paragraph("Led the platform team of 8.")
    p1 = doc.add_paragraph("Scaled API to 100k rps.")
    p1.style = doc.styles["List Paragraph"]
    p2 = doc.add_paragraph("Cut p99 latency by 40%.")
    p2.style = doc.styles["List Paragraph"]
    doc.add_paragraph("Skills: Go, Kubernetes")
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("MIT | BS, Computer Science\n2008 - 2012 | Cambridge, MA")
    doc.add_paragraph("Stanford | MS, Computer Science\n2012 - 2014 | Stanford, CA")
    doc.add_paragraph("SKILLS")
    sk = doc.add_paragraph("Languages: Go, Python")
    sk.style = doc.styles["List Paragraph"]

    path = tmp_path / "synthetic.docx"
    doc.save(str(path))
    return path


def test_synthetic_resume_with_pipe_separator(tmp_path: Path) -> None:
    path = _make_synthetic_resume(tmp_path)
    r = parse_resume(path)
    assert r.header.name == "Alice Example"
    assert r.header.title == "STAFF ENGINEER"
    assert len(r.header.contact_lines) == 1
    assert len(r.experience) == 1
    role = r.experience[0]
    assert role.company == "AcmeCorp"
    assert role.title == "Staff Engineer"
    assert role.dates == "Jan 2020 - Present"
    assert role.location == "Remote"
    assert role.bullets == ["Scaled API to 100k rps.", "Cut p99 latency by 40%."]
    assert role.skills_line == "Skills: Go, Kubernetes"
    assert len(r.education) == 2
    assert r.education[0].institution == "MIT"
    assert r.education[1].institution == "Stanford"


# ---------- error paths ----------


def test_missing_section_raises(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Name")
    doc.add_paragraph("TITLE")
    doc.add_paragraph("Some summary.")
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("Company - Title")
    path = tmp_path / "no_edu.docx"
    doc.save(str(path))
    with pytest.raises(ResumeParseError, match="education"):
        parse_resume(path)


def test_empty_docx_raises(tmp_path: Path) -> None:
    doc = Document()
    path = tmp_path / "empty.docx"
    doc.save(str(path))
    with pytest.raises(ResumeParseError):
        parse_resume(path)
