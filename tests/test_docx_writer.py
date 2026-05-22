"""Tests for src.docx.writer.

Covers:
  - Round-trip identity: parse -> write -> re-parse yields the same Resume.
  - Per-section edits land in the right paragraph indices.
  - Run-level formatting (bold / italic / font size) on untouched
    paragraphs and on edited paragraphs (first-run formatting wins).
  - Skills section growth: surplus categories merge into the last line.
  - Defensive guards.
"""

from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from src.docx.writer import write_resume
from src.stages.resume_parser import parse_resume

SAMPLE_RESUME = Path(__file__).parent.parent / "profiles" / "sample" / "Katharine Berry.docx"


# ---------- round-trip identity ----------


@pytest.mark.skipif(not SAMPLE_RESUME.exists(), reason="sample resume missing")
def test_round_trip_identity_against_real_resume(tmp_path: Path) -> None:
    original = parse_resume(SAMPLE_RESUME)
    out_path = tmp_path / "rt.docx"
    write_resume(SAMPLE_RESUME, original, out_path)

    reparsed = parse_resume(out_path)
    assert reparsed.header.name == original.header.name
    assert reparsed.header.title == original.header.title
    assert reparsed.summary.text == original.summary.text
    assert [e.company for e in reparsed.experience] == [e.company for e in original.experience]
    assert [e.dates for e in reparsed.experience] == [e.dates for e in original.experience]
    for orig_role, new_role in zip(original.experience, reparsed.experience):
        assert orig_role.bullets == new_role.bullets
        assert orig_role.skills_line == new_role.skills_line
    assert reparsed.education[0].institution == original.education[0].institution
    assert reparsed.skills_section.categories == original.skills_section.categories


@pytest.mark.skipif(not SAMPLE_RESUME.exists(), reason="sample resume missing")
def test_paragraph_count_stable_after_write(tmp_path: Path) -> None:
    original = parse_resume(SAMPLE_RESUME)
    out_path = tmp_path / "rt.docx"
    write_resume(SAMPLE_RESUME, original, out_path)

    src_doc = Document(str(SAMPLE_RESUME))
    out_doc = Document(str(out_path))
    assert len(src_doc.paragraphs) == len(out_doc.paragraphs)


# ---------- targeted edits ----------


@pytest.mark.skipif(not SAMPLE_RESUME.exists(), reason="sample resume missing")
def test_modified_summary_round_trips(tmp_path: Path) -> None:
    original = parse_resume(SAMPLE_RESUME)
    edited = original.model_copy(deep=True)
    edited.summary.text = (
        "Senior engineer with 11+ years building distributed developer tools at Google and Pebble."
    )

    out_path = tmp_path / "edit_summary.docx"
    write_resume(SAMPLE_RESUME, edited, out_path)
    reparsed = parse_resume(out_path)
    assert reparsed.summary.text == edited.summary.text
    # Other fields untouched.
    assert reparsed.experience[0].company == original.experience[0].company


@pytest.mark.skipif(not SAMPLE_RESUME.exists(), reason="sample resume missing")
def test_modified_bullets_round_trip(tmp_path: Path) -> None:
    original = parse_resume(SAMPLE_RESUME)
    edited = original.model_copy(deep=True)
    edited.experience[0].bullets = [
        b.replace("developer tools", "distributed developer tools") for b in edited.experience[0].bullets
    ]
    out_path = tmp_path / "edit_bullets.docx"
    write_resume(SAMPLE_RESUME, edited, out_path)
    reparsed = parse_resume(out_path)
    assert reparsed.experience[0].bullets == edited.experience[0].bullets


@pytest.mark.skipif(not SAMPLE_RESUME.exists(), reason="sample resume missing")
def test_modified_title_with_suffix_round_trips(tmp_path: Path) -> None:
    original = parse_resume(SAMPLE_RESUME)
    edited = original.model_copy(deep=True)
    edited.header.title = "SENIOR SOFTWARE ENGINEER, BACKEND PLATFORMS"
    out_path = tmp_path / "edit_title.docx"
    write_resume(SAMPLE_RESUME, edited, out_path)
    reparsed = parse_resume(out_path)
    assert reparsed.header.title == "SENIOR SOFTWARE ENGINEER, BACKEND PLATFORMS"


@pytest.mark.skipif(not SAMPLE_RESUME.exists(), reason="sample resume missing")
def test_modified_skills_section_round_trips(tmp_path: Path) -> None:
    original = parse_resume(SAMPLE_RESUME)
    edited = original.model_copy(deep=True)
    # Append an addition to the "Other" category, which already exists.
    edited.skills_section.categories["Other"] = edited.skills_section.categories["Other"] + ["Real-Time Pipelines"]
    edited.skills_section.raw_lines = [
        f"{cat}: " + ", ".join(items) for cat, items in edited.skills_section.categories.items()
    ]
    out_path = tmp_path / "edit_skills.docx"
    write_resume(SAMPLE_RESUME, edited, out_path)
    reparsed = parse_resume(out_path)
    assert "Real-Time Pipelines" in reparsed.skills_section.categories["Other"]


# ---------- formatting preservation (synthetic doc with bold/italic) ----------


def _make_formatted_resume_docx(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Alice Bold")
    doc.add_paragraph("STAFF ENGINEER")
    doc.add_paragraph("NYC | alice@example.com | (555) 010-2020 | linkedin.com/in/alice")
    doc.add_paragraph("")
    doc.add_paragraph("Staff engineer with 10 years building distributed systems.")
    doc.add_paragraph("PROFESSIONAL EXPERIENCE")
    doc.add_paragraph("AcmeCorp | Staff Engineer")
    doc.add_paragraph("Jan 2020 - Present | Remote")

    # Intro paragraph with formatted runs: "Led" (bold) + " a team of 8 engineers." (plain)
    intro = doc.add_paragraph()
    r1 = intro.add_run("Led")
    r1.bold = True
    intro.add_run(" a team of 8 engineers.")

    # Bullet with italic + bold mixed: "Scaled API to 100k rps." with "100k" in bold italic
    bullet = doc.add_paragraph()
    bullet.style = doc.styles["List Paragraph"]
    rb1 = bullet.add_run("Scaled API to ")
    rb1.italic = True
    rb2 = bullet.add_run("100k rps")
    rb2.bold = True
    rb2.italic = True
    rb3 = bullet.add_run(".")
    rb3.italic = True

    bullet2 = doc.add_paragraph()
    bullet2.style = doc.styles["List Paragraph"]
    bullet2.add_run("Mentored 6 engineers.")

    doc.add_paragraph("Skills: Go, Kubernetes")
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("MIT | BS, Computer Science\n2008 - 2012 | Cambridge, MA")
    doc.add_paragraph("SKILLS")
    skp = doc.add_paragraph("Languages: Go, Python")
    skp.style = doc.styles["List Paragraph"]

    path = tmp_path / "formatted.docx"
    doc.save(str(path))
    return path


def test_first_run_formatting_preserved_when_paragraph_is_edited(tmp_path: Path) -> None:
    path = _make_formatted_resume_docx(tmp_path)
    resume = parse_resume(path)
    edited = resume.model_copy(deep=True)
    edited.experience[0].intro = "Led a team of 8 engineers building real-time systems."

    out_path = tmp_path / "fmt_out.docx"
    write_resume(path, edited, out_path)

    out_doc = Document(str(out_path))
    intro_idx = resume.experience[0].indices.intro_idx
    assert intro_idx is not None
    intro_paragraph = out_doc.paragraphs[intro_idx]
    runs = intro_paragraph.runs
    # Run 0 holds the entire new text and keeps its original bold attribute.
    assert runs[0].text == "Led a team of 8 engineers building real-time systems."
    assert runs[0].bold is True
    # Subsequent runs are emptied but the run elements remain (their original
    # formatting still hangs off the XML — irrelevant for rendering empty runs).
    for r in runs[1:]:
        assert r.text == ""


def test_untouched_paragraph_keeps_all_run_formatting(tmp_path: Path) -> None:
    path = _make_formatted_resume_docx(tmp_path)
    resume = parse_resume(path)
    # Edit ONLY the summary; do not touch the formatted bullet.
    edited = resume.model_copy(deep=True)
    edited.summary.text = "Staff engineer with 10 years of distributed-systems experience."

    out_path = tmp_path / "untouched.docx"
    write_resume(path, edited, out_path)

    out_doc = Document(str(out_path))
    # The first bullet is the one with mixed italic + bold formatting.
    bullet_idx = resume.experience[0].indices.bullet_idxs[0]
    src_doc = Document(str(path))
    src_runs = src_doc.paragraphs[bullet_idx].runs
    out_runs = out_doc.paragraphs[bullet_idx].runs

    assert [r.text for r in out_runs] == [r.text for r in src_runs]
    assert [r.bold for r in out_runs] == [r.bold for r in src_runs]
    assert [r.italic for r in out_runs] == [r.italic for r in src_runs]


def test_paragraph_style_preserved_on_edit(tmp_path: Path) -> None:
    path = _make_formatted_resume_docx(tmp_path)
    resume = parse_resume(path)
    edited = resume.model_copy(deep=True)
    edited.experience[0].bullets[1] = "Mentored 6 engineers across three teams."

    out_path = tmp_path / "style.docx"
    write_resume(path, edited, out_path)

    out_doc = Document(str(out_path))
    bullet_idx = resume.experience[0].indices.bullet_idxs[1]
    assert out_doc.paragraphs[bullet_idx].style.name == "List Paragraph"


# ---------- skills section overflow ----------


def test_skills_section_surplus_lines_merge_into_last_paragraph(tmp_path: Path) -> None:
    path = _make_formatted_resume_docx(tmp_path)
    resume = parse_resume(path)
    edited = resume.model_copy(deep=True)
    # Source has 1 category ("Languages"); add a fake "Other" with one item.
    edited.skills_section.categories["Other"] = ["Real-Time Pipelines"]
    edited.skills_section.raw_lines = [
        "Languages: Go, Python",
        "Other: Real-Time Pipelines",
    ]
    out_path = tmp_path / "skills_overflow.docx"
    write_resume(path, edited, out_path)

    out_doc = Document(str(out_path))
    # Only one content paragraph existed in the source; surplus items should
    # be merged into it.
    content_idx = resume.skills_section.indices.content_idxs[-1]
    text = out_doc.paragraphs[content_idx].text
    assert "Languages: Go, Python" in text
    assert "Real-Time Pipelines" in text


# ---------- new bullet insertion + defensive guards (v2 hybrid pipeline) ----------


def test_new_bullets_inserted_before_skills_line(tmp_path: Path) -> None:
    path = _make_formatted_resume_docx(tmp_path)
    resume = parse_resume(path)
    edited = resume.model_copy(deep=True)
    # Append a new bullet beyond the original bullet_idxs count.
    edited.experience[0].bullets = edited.experience[0].bullets + ["Built a new Kafka pipeline."]
    out_path = tmp_path / "new_bullets.docx"
    write_resume(path, edited, out_path)

    out_doc = Document(str(out_path))
    src_doc = Document(str(path))
    # New paragraph added → output count = source + 1.
    assert len(out_doc.paragraphs) == len(src_doc.paragraphs) + 1

    # The inserted paragraph carries the bullet text.
    inserted_texts = [p.text for p in out_doc.paragraphs]
    assert "Built a new Kafka pipeline." in inserted_texts


def test_new_bullets_carry_list_paragraph_style(tmp_path: Path) -> None:
    path = _make_formatted_resume_docx(tmp_path)
    resume = parse_resume(path)
    edited = resume.model_copy(deep=True)
    edited.experience[0].bullets = edited.experience[0].bullets + ["A new JD-aligned achievement."]
    out_path = tmp_path / "new_bullet_style.docx"
    write_resume(path, edited, out_path)

    out_doc = Document(str(out_path))
    for p in out_doc.paragraphs:
        if p.text == "A new JD-aligned achievement.":
            assert p.style.name == "List Paragraph"
            return
    pytest.fail("inserted bullet not found in output")


def test_existing_bullets_unchanged_when_new_bullet_inserted(tmp_path: Path) -> None:
    path = _make_formatted_resume_docx(tmp_path)
    resume = parse_resume(path)
    original_bullets = list(resume.experience[0].bullets)

    edited = resume.model_copy(deep=True)
    edited.experience[0].bullets = edited.experience[0].bullets + ["Extra bullet text."]
    out_path = tmp_path / "extra.docx"
    write_resume(path, edited, out_path)

    # Reparse and confirm original bullets are still there in order.
    reparsed = parse_resume(out_path)
    for orig_bullet in original_bullets:
        assert orig_bullet in reparsed.experience[0].bullets


def test_dropping_existing_bullets_still_raises(tmp_path: Path) -> None:
    """The writer still rejects fewer-bullets-than-indices as a defensive guard."""
    path = _make_formatted_resume_docx(tmp_path)
    resume = parse_resume(path)
    edited = resume.model_copy(deep=True)
    edited.experience[0].bullets = edited.experience[0].bullets[:1]  # drop one
    with pytest.raises(ValueError, match="bullet count mismatch"):
        write_resume(path, edited, tmp_path / "should_not_save.docx")


def test_newlines_in_replacement_text_are_collapsed_to_spaces(tmp_path: Path) -> None:
    path = _make_formatted_resume_docx(tmp_path)
    resume = parse_resume(path)
    edited = resume.model_copy(deep=True)
    edited.summary.text = "Line one.\nLine two."
    out_path = tmp_path / "newline.docx"
    write_resume(path, edited, out_path)

    out_doc = Document(str(out_path))
    summary_idx = resume.summary.paragraph_idxs[0]
    assert "\n" not in out_doc.paragraphs[summary_idx].text
    assert "Line one. Line two." in out_doc.paragraphs[summary_idx].text
